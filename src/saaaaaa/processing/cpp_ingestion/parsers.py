"""
Document parsers for various formats (PDF, DOCX, HTML).

Implements deterministic parsing with stable output using production libraries.
"""

import io
from typing import Any, Dict, List

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


class DocumentParser:
    """Multi-format document parser with production-ready implementations."""
    
    def parse(self, binary_data: bytes, mime_type: str) -> Dict[str, Any]:
        """
        Parse document based on MIME type.
        
        Args:
            binary_data: Raw document bytes
            mime_type: MIME type of document
            
        Returns:
            Dictionary with parsed structure including pages, text, tables, and metadata
        """
        if mime_type == "application/pdf":
            return self._parse_pdf(binary_data)
        elif "wordprocessingml" in mime_type:
            return self._parse_docx(binary_data)
        elif mime_type == "text/html":
            return self._parse_html(binary_data)
        else:
            raise ValueError(f"Unsupported MIME type: {mime_type}")
    
    def _parse_pdf(self, binary_data: bytes) -> Dict[str, Any]:
        """
        Parse PDF document using pdfplumber.
        
        Extracts text, tables, and structural information from each page.
        """
        if pdfplumber is None:
            raise RuntimeError("pdfplumber is required for PDF parsing. Install with: pip install pdfplumber")
        
        pages = []
        all_tables = []
        
        try:
            # Use pdfplumber for comprehensive PDF parsing
            with pdfplumber.open(io.BytesIO(binary_data)) as pdf:
                metadata = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text with layout preservation
                    text = page.extract_text() or ""
                    
                    # Extract tables from page
                    page_tables = page.extract_tables() or []
                    
                    # Convert tables to structured format
                    structured_tables = []
                    for table_idx, table in enumerate(page_tables):
                        if table:
                            structured_table = {
                                "table_id": f"page{page_num}_table{table_idx}",
                                "page": page_num,
                                "rows": table,
                                "headers": table[0] if table else [],
                                "data_rows": table[1:] if len(table) > 1 else [],
                            }
                            structured_tables.append(structured_table)
                            all_tables.append(structured_table)
                    
                    # Get page dimensions and objects
                    page_info = {
                        "page_num": page_num,
                        "text": text,
                        "tables": structured_tables,
                        "objects": [],
                        "width": page.width,
                        "height": page.height,
                        "chars": len(page.chars),
                    }
                    
                    pages.append(page_info)
            
            return {
                "format": "pdf",
                "pages": pages,
                "tables": all_tables,
                "metadata": metadata,
                "page_count": len(pages),
            }
            
        except Exception as e:
            raise RuntimeError(f"PDF parsing failed: {e}")
    
    def _parse_docx(self, binary_data: bytes) -> Dict[str, Any]:
        """
        Parse DOCX document using python-docx.
        
        Extracts paragraphs, tables, and structural information.
        """
        if DocxDocument is None:
            raise RuntimeError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        pages = []
        all_tables = []
        
        try:
            # Use python-docx for DOCX parsing
            doc = DocxDocument(io.BytesIO(binary_data))
            
            # Extract core properties
            metadata = {}
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }
            
            # Collect all text (paragraphs)
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_text.append(para.text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                structured_table = {
                    "table_id": f"docx_table{table_idx}",
                    "page": 1,  # DOCX doesn't have explicit pages
                    "rows": [],
                    "headers": [],
                    "data_rows": [],
                }
                
                # Extract table data
                for row_idx, row in enumerate(table.rows):
                    row_data = [cell.text.strip() for cell in row.cells]
                    structured_table["rows"].append(row_data)
                    
                    if row_idx == 0:
                        structured_table["headers"] = row_data
                    else:
                        structured_table["data_rows"].append(row_data)
                
                all_tables.append(structured_table)
            
            # Create single page (DOCX is continuous)
            page_info = {
                "page_num": 1,
                "text": "\n".join(all_text),
                "tables": all_tables,
                "objects": [],
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }
            
            pages.append(page_info)
            
            return {
                "format": "docx",
                "pages": pages,
                "tables": all_tables,
                "metadata": metadata,
                "page_count": 1,
            }
            
        except Exception as e:
            raise RuntimeError(f"DOCX parsing failed: {e}")
    
    def _parse_html(self, binary_data: bytes) -> Dict[str, Any]:
        """
        Parse HTML document using BeautifulSoup.
        
        Extracts text, tables, and structural elements.
        """
        if BeautifulSoup is None:
            raise RuntimeError("beautifulsoup4 is required for HTML parsing. Install with: pip install beautifulsoup4 lxml")
        
        pages = []
        all_tables = []
        
        try:
            # Decode HTML
            html_text = binary_data.decode("utf-8", errors="ignore")
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(html_text, "lxml")
            
            # Extract title
            title = soup.title.string if soup.title else ""
            
            # Extract main text content (excluding scripts and styles)
            for script in soup(["script", "style"]):
                script.decompose()
            
            text = soup.get_text(separator="\n", strip=True)
            
            # Extract tables
            html_tables = soup.find_all("table")
            for table_idx, table in enumerate(html_tables):
                rows = []
                headers = []
                
                # Extract headers
                header_row = table.find("thead")
                if header_row:
                    headers = [th.get_text(strip=True) for th in header_row.find_all(["th", "td"])]
                
                # Extract all rows
                for tr in table.find_all("tr"):
                    cells = [td.get_text(strip=True) for td in tr.find_all(["td", "th"])]
                    if cells:
                        rows.append(cells)
                
                if rows:
                    structured_table = {
                        "table_id": f"html_table{table_idx}",
                        "page": 1,
                        "rows": rows,
                        "headers": headers or (rows[0] if rows else []),
                        "data_rows": rows[1:] if len(rows) > 1 else rows,
                    }
                    all_tables.append(structured_table)
            
            # Create single page
            page_info = {
                "page_num": 1,
                "text": text,
                "tables": all_tables,
                "objects": [],
                "title": title,
                "table_count": len(all_tables),
            }
            
            pages.append(page_info)
            
            return {
                "format": "html",
                "pages": pages,
                "tables": all_tables,
                "metadata": {"title": title},
                "page_count": 1,
            }
            
        except Exception as e:
            raise RuntimeError(f"HTML parsing failed: {e}")
